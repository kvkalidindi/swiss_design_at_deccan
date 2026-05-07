"""Build deccan.dotx using python-docx.

Word template with the four pieces of brand "furniture":
  1. Cover (face) page  — logo, title, metadata block
  2. Running header     — small logo + doc title (every body page)
  3. Running footer     — confidentiality + page number (every body page)
  4. End page           — centered logo + brand line

The cover and end page do not show the running header/footer; the cover uses
"different first page", the end page sits in its own section.
See skill/references/document-furniture.md for the canonical spec.
"""
from __future__ import annotations

import json
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor as DocxRGB
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from scripts.lib.office_theme import build_theme_xml

# Stone-100 hex for code chip / block backgrounds (no leading #).
STONE_100 = "F5F5F4"
MONO_FONT = "IBM Plex Mono"

ROOT = Path(__file__).resolve().parents[2]
PALETTE = ROOT / "outputs" / "palette.json"
LOGO = ROOT / "data" / "logo.png"
OUT = ROOT / "office" / "templates" / "deccan.dotx"

# Sample metadata used in the template's cover page.
META_LABELS = ["DOCUMENT TYPE", "PREPARED BY", "DATE", "VERSION", "CLASSIFICATION"]
META_VALUES = ["Report", "Author Name", "YYYY-MM-DD", "v1.0", "Confidential"]


def _palette() -> dict:
    return json.loads(PALETTE.read_text(encoding="utf-8"))


def _add_page_number_field(paragraph) -> None:
    """Insert a bare PAGE field (just the number) into the given paragraph."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE   \\* MERGEFORMAT"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = paragraph.add_run()
    page_run.font.name = "IBM Plex Sans"
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = DocxRGB.from_string("595959")
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_text)
    page_run._r.append(fld_end)


def _enable_update_fields_on_open(doc) -> None:
    """Add <w:updateFields w:val="true"/> to settings so PAGE renders on first open."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def _set_normal_style_full_width(doc) -> None:
    """Ensure the Normal paragraph style fills the content area — no right indent,
    no max width — so body text takes the full live width per the >=80% rule."""
    if "Normal" not in doc.styles:
        return
    s = doc.styles["Normal"]
    pf = s.paragraph_format
    pf.left_indent = Inches(0)
    pf.right_indent = Inches(0)
    pf.first_line_indent = Inches(0)
    s.font.name = "IBM Plex Sans"
    s.font.size = Pt(11)


def _add_shading(element, fill_hex: str) -> None:
    """Append <w:shd w:val="clear" w:color="auto" w:fill="..."/> to a pPr or rPr."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def _register_code_styles(doc) -> None:
    """Register Code Inline (character) and Code Block (paragraph) styles with
    IBM Plex Mono + stone-100 shading per skill/references/document-furniture.md."""
    styles = doc.styles

    # Code Inline — character style for runs of inline code/identifiers.
    if "Code Inline" not in styles:
        ci = styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
        ci.font.name = MONO_FONT
        ci.font.size = Pt(10)
        ci.font.color.rgb = DocxRGB.from_string("1C1917")
        rPr = ci.element.get_or_add_rPr()
        _add_shading(rPr, STONE_100)
        # Force the East-Asian / complex-script font slot to the mono face too.
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), MONO_FONT)

    # Code Block — paragraph style for multi-line code blocks.
    if "Code Block" not in styles:
        cb = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        cb.font.name = MONO_FONT
        cb.font.size = Pt(10)
        cb.font.color.rgb = DocxRGB.from_string("1C1917")
        pf = cb.paragraph_format
        pf.left_indent = Inches(0)
        pf.right_indent = Inches(0)
        pf.first_line_indent = Inches(0)
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        pf.keep_together = True
        rPr = cb.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), MONO_FONT)
        # Shade the paragraph (block fill).
        pPr = cb.element.get_or_add_pPr()
        _add_shading(pPr, STONE_100)


def _set_paragraph_bottom_border(paragraph, color_hex: str) -> None:
    """Add a bottom border to a paragraph (used for header/footer rules)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # 0.5pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_top_border(paragraph, color_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def _build_running_header(section, title_text: str) -> None:
    """Write the running header into `section.header`.

    Layout: [logo] .................... [doc title], with a thin bottom rule.
    """
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_TAB_ALIGNMENT.RIGHT,
    )
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(0.6))
    p.add_run("\t")
    title_run = p.add_run(title_text)
    title_run.font.name = "IBM Plex Sans"
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = DocxRGB.from_string("595959")  # stone-700-ish gray
    _set_paragraph_bottom_border(p, "D6D3D1")  # stone-200


def _build_running_footer(section, classification: str = "Confidential") -> None:
    """Write the running footer into `section.footer`.

    Layout: [Deccan Chemicals · classification] .... [Page X of Y], thin top rule.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_TAB_ALIGNMENT.RIGHT,
    )
    left = p.add_run(f"Deccan Chemicals · {classification}")
    left.font.name = "IBM Plex Sans"
    left.font.size = Pt(9)
    left.font.color.rgb = DocxRGB.from_string("595959")
    p.add_run("\t")
    _add_page_number_field(p)
    _set_paragraph_top_border(p, "D6D3D1")


def _disable_section_furniture(section) -> None:
    """Suppress header/footer on a section (used for end-page section)."""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    # Clear by writing a single empty paragraph
    for para in list(section.header.paragraphs):
        for run in para.runs:
            run.text = ""
    for para in list(section.footer.paragraphs):
        for run in para.runs:
            run.text = ""


def _build_cover(doc, blue_500: str) -> None:
    """Write the cover (face) page composition into the first page of `doc`."""
    # Top whitespace
    for _ in range(3):
        doc.add_paragraph()

    if LOGO.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        logo_p.add_run().add_picture(str(LOGO), width=Inches(2.5))

    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title_p.add_run("Document Title")
    title_run.font.name = "IBM Plex Sans"
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = DocxRGB.from_string(blue_500)
    title_run.font.bold = False  # font-light per design system

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sub_run = sub_p.add_run("Subtitle / one-line summary")
    sub_run.font.name = "IBM Plex Sans"
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = DocxRGB.from_string("595959")

    # Accent rule below subtitle
    rule_p = doc.add_paragraph()
    _set_paragraph_bottom_border(rule_p, blue_500)

    doc.add_paragraph()

    # Metadata block (caption-style key/value pairs)
    for label, value in zip(META_LABELS, META_VALUES):
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(2)
        label_run = meta_p.add_run(f"{label}    ")
        label_run.font.name = "IBM Plex Sans"
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = DocxRGB.from_string("8A8786")
        label_run.font.bold = False
        value_run = meta_p.add_run(value)
        value_run.font.name = "IBM Plex Sans"
        value_run.font.size = Pt(10)
        value_run.font.color.rgb = DocxRGB.from_string("1C1917")


def _build_end_page(doc, blue_500: str) -> None:
    """Write the end-page composition (centered logo + brand line)."""
    # Vertical breathing room above
    for _ in range(10):
        doc.add_paragraph()

    if LOGO.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logo_p.add_run().add_picture(str(LOGO), width=Inches(1.8))

    doc.add_paragraph()

    brand_p = doc.add_paragraph()
    brand_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    brand_run = brand_p.add_run("Deccan Chemicals")
    brand_run.font.name = "IBM Plex Sans"
    brand_run.font.size = Pt(14)
    brand_run.font.color.rgb = DocxRGB.from_string("1C1917")

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_run = contact_p.add_run("deccanchemicals.com · Hyderabad, India")
    contact_run.font.name = "IBM Plex Sans"
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = DocxRGB.from_string("8A8786")


def _set_section_different_first_page(section) -> None:
    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)


def _configure_heading_styles(doc, blue_500: str) -> None:
    """Set heading styles to enforce Swiss Deccan flow rules.

    - keep_with_next: heading sticks to its first body paragraph (no orphan headings).
    - keep_lines_together: heading itself never breaks across pages.
    - page_break_before on H1: every top-level section starts on a new page,
      which trivially satisfies the "no new section in the bottom 25%" rule.
    - generous space_before on H2 to give visual separation.
    """
    styles = doc.styles
    for level, page_break in [("Heading 1", True), ("Heading 2", False), ("Heading 3", False)]:
        if level not in styles:
            continue
        s = styles[level]
        s.font.name = "IBM Plex Sans"
        s.font.color.rgb = DocxRGB.from_string(blue_500)
        s.font.bold = False
        pf = s.paragraph_format
        pf.keep_with_next = True
        pf.keep_together = True
        pf.widow_control = True
        if page_break:
            pf.page_break_before = True
        if level == "Heading 1":
            pf.space_before = Pt(24)
            pf.space_after = Pt(12)
        elif level == "Heading 2":
            pf.space_before = Pt(18)
            pf.space_after = Pt(6)


def emit_dotx() -> Path:
    palette = _palette()
    blue_500 = palette["blue"]["500"]["hex"].lstrip("#")

    doc = Document()

    # Section 0 — cover + body. Cover is the first page; running furniture
    # appears from page 2 onward via "different first page".
    section = doc.sections[0]
    # Margins of 0.8" on Letter give 6.9" / 8.5" = 81.2% live content width,
    # satisfying the >=80% rule (see skill/references/document-furniture.md).
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    _set_section_different_first_page(section)
    _configure_heading_styles(doc, blue_500)
    _set_normal_style_full_width(doc)
    _register_code_styles(doc)
    _enable_update_fields_on_open(doc)

    # First-page header/footer (suppressed: Word renders them blank because
    # different-first-page is on, but we still need empty placeholders).
    fp_header = section.first_page_header
    fp_header.is_linked_to_previous = False
    fp_footer = section.first_page_footer
    fp_footer.is_linked_to_previous = False

    # Running header/footer for body pages
    _build_running_header(section, "Document Title")
    _build_running_footer(section, "Confidential")

    # Cover page composition
    _build_cover(doc, blue_500)

    # Page break out of cover into body content
    doc.add_page_break()

    # --- Body: example styled content ---
    h1 = doc.add_heading("Heading 1 example", level=1)
    for run in h1.runs:
        run.font.name = "IBM Plex Sans"
        run.font.color.rgb = DocxRGB.from_string(blue_500)

    h2 = doc.add_heading("Heading 2 example", level=2)
    for run in h2.runs:
        run.font.name = "IBM Plex Sans"

    p = doc.add_paragraph()
    for run in p.runs:
        run.font.name = "IBM Plex Sans"
        run.font.size = Pt(11)
    p.add_run(
        "Body paragraph using the Deccan typography stack. IBM Plex Sans at 11pt "
        "with 1.5 line height. Use opacity (not different colors) for text "
        "hierarchy per the Deccan design system. Reference inline code like "
    )
    chip = p.add_run("emit_dotx()")
    chip.style = doc.styles["Code Inline"]
    p.add_run(" or a file path like ")
    chip2 = p.add_run("scripts/lib/office_docx.py")
    chip2.style = doc.styles["Code Inline"]
    p.add_run(" using the Code Inline character style.")
    for run in p.runs:
        if run.style is None or run.style.name == "Default Paragraph Font":
            run.font.name = "IBM Plex Sans"
            run.font.size = Pt(11)

    # Code Block example
    cb = doc.add_paragraph(style="Code Block")
    cb.add_run("python -m scripts._09_emit_office\nWrote office/templates/deccan.dotx (49,972 bytes)")

    # Explicit page break out of body content before the end page.
    doc.add_page_break()

    # --- End page section: own section, no running header/footer, no page # ---
    end_section = doc.add_section(WD_SECTION.NEW_PAGE)
    end_section.top_margin = Inches(0.8)
    end_section.bottom_margin = Inches(0.8)
    end_section.left_margin = Inches(0.8)
    end_section.right_margin = Inches(0.8)
    _disable_section_furniture(end_section)
    _build_end_page(doc, blue_500)

    tmp = OUT.with_suffix(".docx")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(tmp)

    _replace_theme_in_docx(tmp, build_theme_xml(palette))
    _convert_to_dotx(tmp, OUT)
    tmp.unlink()
    return OUT


def _replace_theme_in_docx(path: Path, theme_xml: str) -> None:
    """Replace word/theme/theme1.xml in the .docx zip with our theme."""
    buf = BytesIO()
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            theme_replaced = False
            for item in zin.namelist():
                data = zin.read(item)
                if item == "word/theme/theme1.xml":
                    data = theme_xml.encode("utf-8")
                    theme_replaced = True
                zout.writestr(item, data)
            if not theme_replaced:
                # python-docx older versions sometimes omit the theme file;
                # add it if missing.
                zout.writestr("word/theme/theme1.xml", theme_xml)
    path.write_bytes(buf.getvalue())


def _convert_to_dotx(src: Path, dst: Path) -> None:
    """Convert .docx package to .dotx by updating [Content_Types].xml."""
    buf = BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    text = text.replace(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                    )
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    dst.write_bytes(buf.getvalue())
